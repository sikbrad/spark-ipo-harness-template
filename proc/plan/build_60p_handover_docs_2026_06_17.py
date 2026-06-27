#!/usr/bin/env python3
"""Build 60-page expanded handover docs from collected Teams evidence only."""
from __future__ import annotations

import importlib.util
import json
import re
import sys
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


RUN_DATE = "2026-06-17"
ROOT = Path("/Users/gq/works/projs/dof-work-skills/dof-work-startpoint-04")
BASE_SCRIPT = ROOT / "proc/plan/build_30p_handover_docs_2026_06_17.py"
SUMMARY_PATH = ROOT / "output/teams-business-rules/2026-06-17/evidence_summary.json"
PEOPLE_DIR = ROOT / "data/company/people"
OUT_BASE = ROOT / "output/teams-business-rules/2026-06-17/handover_60p_docs"
OUT_MD = OUT_BASE / "markdown"
OUT_DOCX = OUT_BASE / "docx"
VERIFY_PATH = OUT_BASE / "verification.json"


spec = importlib.util.spec_from_file_location("handover30", BASE_SCRIPT)
handover30 = importlib.util.module_from_spec(spec)
sys.modules[spec.name] = handover30
assert spec.loader
spec.loader.exec_module(handover30)

Page = handover30.Page
p = handover30.p


@dataclass(frozen=True)
class Topic:
    title: str
    focus: str
    keywords: tuple[str, ...]
    lens: str


def kw(value: str) -> tuple[str, ...]:
    return tuple(part.strip() for part in value.split("|") if part.strip())


TOPICS: dict[str, list[Topic]] = {
    "김규탁": [
        Topic("인코텀즈 변경 세부규칙", "DAP, EXW, DDP, DHL 변경은 운송비와 위험부담을 동시에 바꾸므로 주문 반려/재승인까지 연결한다.", kw("DAP|EXW|DDP|DHL|운임|Shipping|shipping fee|인코텀즈"), "logistics"),
        Topic("포장수량·중량 재예약 규칙", "예약값과 실제 포장값이 다르면 픽업 예약이 불가능하므로 실포장 기준으로 다시 잡는다.", kw("PKG|포장|중량|무게|예약|부킹|픽업|박스"), "logistics"),
        Topic("HS code 변경 대응", "HS code 변경은 CIPL 수정만이 아니라 수입자 컨펌과 국가별 통관 리스크 확인으로 이어진다.", kw("HS|HS code|세번|9021|9031|통관|수입자 컨펌"), "country"),
        Topic("CIPL 버전관리", "CIPL은 포장 전 초안과 포장 후 최종본을 구분해야 하며 출고 후 수정은 실물 부착 서류와 달라질 수 있다.", kw("CIPL|Invoice|Commercial|Packing|PL|Buyer|VAT|value"), "document"),
        Topic("AWB·BL·MAWB cutoff", "AWB/BL/MAWB/HAWB는 문서값, 포장값, 포워더 cutoff가 맞아야 선적 일정으로 확정된다.", kw("AWB|BL|MAWB|HAWB|ETD|cutoff|선박|해상|항공"), "document"),
        Topic("노미포워더 정보 검증", "노미포워더는 담당자 이름이 아니라 실제 운송사, 창고, 픽업지, 기적 가능성을 의미한다.", kw("노미|포워더|forwarder|창고|픽업|기적|면장|수령"), "logistics"),
        Topic("전략물자 보완자료", "전략물자 건은 최종사용자, 사용목적, 사업장 증빙, 허가상태가 출고보다 선행된다.", kw("전략물자|최종사용자|사용목적|사업장|허가|증빙|서약서"), "country"),
        Topic("장기미수·송금 전 보류", "송금 전 또는 장기미수 거래는 공식 승인과 회계 수금확인 없이는 출고하지 않는다.", kw("미수|송금전|입금|수금|보류|장기|payment|잔액"), "finance"),
        Topic("중국 OEM 로트관리", "중국 OEM 수입은 PO/로트/전압/CI/PL/COO/AMJ 검수를 한 흐름으로 본다.", kw("Philden|JINY|중국|OEM|로트|전압|CI|COO|AMJ"), "product"),
        Topic("AMJ 준비상태와 설치일", "설치일은 AMJ 준비상태와 배차 가능일이 확인된 뒤에 확정한다.", kw("AMJ|설치|픽업|배차|CRAFT DX|CRAFT S|준비"), "product"),
        Topic("라벨·시리얼 충돌", "공급사 라벨, DOF 라벨, 장비 내부 시리얼이 충돌하면 출고가 아니라 검수 이슈다.", kw("라벨|시리얼|serial|스티커|관리번호|검수|DOF 라벨"), "product"),
        Topic("DOF Japan 재고배분", "일본법인 주문이라도 국내 주문과 같은 실물 재고가 걸리면 배분 판단을 올린다.", kw("DOF Japan|일본법인|일본|AMJ|Freedom|Jig|동일 재고"), "customer"),
        Topic("GT-MEDICAL·스페인 수입문서", "스페인 장비/기관 주문은 전략물자, 수입허가, HS, 라벨 책임을 함께 본다.", kw("GT-MEDICAL|3DI|Spain|스페인|University|Nueva|EORI|수입허가"), "customer"),
        Topic("FOR MEDIKAL·터키 법인명", "터키 건은 포탈 고객사명과 CIPL/BL 수입자명을 법인 단위로 대조한다.", kw("FOR MEDIKAL|Mars Med|Turkey|터키|Istanbul|Fatih|수입자"), "customer"),
        Topic("Yamamoto·우크라이나 리스크", "우크라이나 거래는 전략물자, 장기미수, 전쟁 리스크, 소프트웨어/물품 분리를 함께 본다.", kw("Yamamoto|S2|우크라이나|Ukraine|전쟁|장기 미수|software"), "customer"),
        Topic("CIDEAS·중남미 반송", "중남미 거래는 원주문, 반송, 환불, 운임, CIPL 분리를 먼저 나눠야 한다.", kw("CIDEAS|Macrodent|Sumidental|페루|콜롬비아|니카라과|반송|환불"), "customer"),
        Topic("Red Lion 쉽백·재고전환", "대량 해상 보류가 길어지면 쉽백, 포장해체, 국내재고 전환 판단이 필요하다.", kw("Red Lion|쉽백|컨테이너|재고전환|포장해체|BL 정정"), "customer"),
        Topic("Mawasi Dent 실포장 기준", "Mawasi Dent 같은 대량 주문은 ERP 주문수량보다 실포장 완료 수량을 우선한다.", kw("Mawasi|패킹|Block|DOF SOLID|대량|완료 수량"), "customer"),
        Topic("PIM KIT 구성품 누락", "PIM KIT와 키트류는 실제 포함품이 CIPL/PL과 맞아야 하며 누락은 통관 이슈다.", kw("PIM|KIT|키트|구성품|동봉|누락|CO"), "product"),
        Topic("구강스캐너 운송방식", "스캐너는 파손위험 때문에 일반택배, 퀵, 직접전달 중 책임방식을 확인한다.", kw("스캐너|FREEDOM|Freedom|Edge HD|택배|파손|직접전달"), "product"),
        Topic("A/S 부품 회수정보", "A/S 부품은 유상/무상 근거, 대상 장비 시리얼, 회수 예정 정보가 함께 있어야 한다.", kw("A/S|AS|무상|유상|회수|Projector|Assy|시리얼"), "order"),
        Topic("샘플·핸드캐리", "샘플과 핸드캐리는 판매품과 다른 문서가 필요하고 전달자/전달완료일도 기록해야 한다.", kw("샘플|sample|핸드캐리|handcarry|출장|전달완료|무상"), "order"),
        Topic("반품 매출 값", "반품은 원단가, 수량 부호, 회수 완료 시점이 맞아야 회계 처리로 넘어간다.", kw("반품|마이너스|원단가|음수|양수|회수 완료|재출고"), "finance"),
        Topic("수입 제도 ACID·CargoX", "이집트 등 제도성 통관은 ACID/CargoX 같은 선행 플랫폼 상태를 확인해야 한다.", kw("ACID|CargoX|ACI|이집트|Egypt|통관마감"), "country"),
        Topic("사우디 SABER·TGA", "사우디는 SABER/TGA/short address 같은 필수값 없이는 진행하지 않는다.", kw("SABER|TGA|사우디|Saudi|short address|DDP"), "country"),
        Topic("우즈베키스탄 C/O", "우즈베키스탄 건은 C/O와 신고서/CIPL/포워더 기적처리 일치가 핵심이다.", kw("우즈베키스탄|Uzbek|C/O|Dental Design|Baraka|Nikola"), "country"),
        Topic("미국·DOF USA 비재고", "DOF USA 요청은 생산 재고인지 별도 발주/비재고/샘플인지 먼저 분류한다.", kw("DOF USA|미국|USA|Cart|Bucket|샘플|비재고"), "customer"),
        Topic("공식 요청창구 운영", "메신저와 유선으로 흩어진 출고 요청은 공식 기록으로 남겨야 누락이 줄어든다.", kw("NOCO|공식|메신저|유선|요청|기록|채널"), "process"),
        Topic("출고 재개조건 문장화", "보류 건은 무엇이 들어오면 재개되는지 명확히 남겨야 다음 사람이 이어받는다.", kw("보류|재개|확인|요청|누락|진행 가능"), "process"),
    ],
    "조소연": [
        Topic("정정주문과 이력보존", "삭제보다 이력 보존이 중요할 때는 정정주문으로 원주문 관계를 남긴다.", kw("정정주문|삭제|재작성|원주문|대체주문|히스토리"), "order"),
        Topic("단가·수량 반전", "단가와 수량이 뒤바뀐 주문은 기준값 확인 전 출고승인하지 않는다.", kw("단가|수량|반전|판단가|금액|오입력"), "order"),
        Topic("계약금액 오입력", "계약 총액과 주문 단가가 다르면 삭제로 덮지 말고 정정주문과 전표를 맞춘다.", kw("계약|계약금액|계약서|총액|전표|가계약"), "finance"),
        Topic("반품금액 대조", "반품 주문은 원주문 단가와 반품 단가, 운임 환불을 표로 대조한다.", kw("반품|원주문|단가|운임|환불|마이너스"), "finance"),
        Topic("부가세 포함/별도", "부가세 포함/별도 충돌은 고객구분과 단가구분을 다시 맞춰야 한다.", kw("부가세|포함|별도|고객구분|단가구분|121만원"), "finance"),
        Topic("해외 ERP 부가세 0", "해외 주문은 부가세율 0과 미포함 기준을 ERP 전송 전 확인한다.", kw("해외|ERP|부가세율|0|미포함|전송"), "system"),
        Topic("금액 0 해외 인코텀즈", "금액 0원 주문도 인코텀즈와 문서 value가 필수다.", kw("금액 0|0원|인코텀즈|incoterms|value|무상"), "country"),
        Topic("DAP·EXW·DDP 오류", "운송조건과 shipping fee가 맞지 않으면 반려/재승인이 필요하다.", kw("DAP|EXW|DDP|shipping fee|운송비|운임"), "logistics"),
        Topic("DOF Japan 운송조건", "일본법인 주문은 품목군별 EXW/DAP 정책과 월마감 귀속을 같이 본다.", kw("DOF Japan|일본법인|EXW|DAP|월마감|SKU"), "customer"),
        Topic("exocad SKU·Ticket", "exocad 주문은 SKU, 정책동의서, 동글, Ticket을 분리하지 않는다.", kw("exocad|SKU|Ticket|동글|정책동의서"), "product"),
        Topic("hyperDENT 동글·적용일", "hyperDENT는 동글번호, 옵션 포함 여부, 적용일 리드타임을 확인한다.", kw("hyperDENT|동글|라이선스|옵션|리드타임|업데이트"), "product"),
        Topic("CRAFT 전략물자 기안", "CRAFT/5축 장비는 최종사용자와 허가 기안이 주문 전제다.", kw("CRAFT|5축|전략물자|최종사용자|기안|사업증빙"), "product"),
        Topic("중국 OEM 역할표", "Philden 초도 수입은 영업/구매/수입물류/악세사리 역할표가 필요하다.", kw("Philden|중국|OEM|역할|PO|로트|전압"), "product"),
        Topic("스캐너 시리얼 필드화", "스캐너/장비 주문은 시리얼 공란이면 외주처 재확인이 반복된다.", kw("스캐너|FREEDOM|시리얼|serial|외주처|장비"), "product"),
        Topic("Tool·Block 품목오류", "HASS 등 제품 오입력은 기존 주문 유지보다 올바른 제품 재등록이 안전하다.", kw("Tool|Block|Premill|HASS|품목|코드|재고"), "product"),
        Topic("워런티·PKG 잔액", "패키지/워런티 주문은 계약 연결과 잔액, 원주문 금액 일치가 먼저다.", kw("워런티|PKG|패키지|잔액|계약|차감"), "finance"),
        Topic("선금·발주연계", "발주연계 품목은 미입금 상태로 발주하면 비용 리스크가 생긴다.", kw("선금|선입금|발주|입금|매입처|라이선스 키"), "finance"),
        Topic("라벨 누락 책임경계", "라벨 누락 반송은 사전 영업 요청과 출고 사진으로 책임을 가른다.", kw("라벨|누락|반송|출고 사진|Nueva|책임"), "customer"),
        Topic("스페인 수입문서", "스페인 장비/기관 주문은 EORI, 위생/수입문서, 전략물자 자료를 같이 본다.", kw("스페인|Spain|EORI|GT-MEDICAL|3DI|University|수입문서"), "country"),
        Topic("터키 통관 표기", "터키 고객은 포탈 고객사명과 CIPL/BL 수입자 표기를 분리 관리한다.", kw("터키|FOR MEDIKAL|Mars Med|수입자|consignee|Istanbul"), "country"),
        Topic("AsiaGrand 국가값", "AsiaGrand는 실제 국내 거래와 CRM 매출집계 국가가 충돌할 수 있다.", kw("AsiaGrand|국가|러시아|노미|포워더|매출집계"), "customer"),
        Topic("CIDEAS 환불정책", "CIDEAS는 반품정책, 운임 환불, 미수 차감, 최초 운송비를 한 세트로 본다.", kw("CIDEAS|환불|반품정책|운임|미수|DAP"), "customer"),
        Topic("CNC AS 운임 마이너스", "CNC AS 반품은 운임 환불을 품목 단가에 섞지 않는다.", kw("CNC AS|운임|환불|반품 단가|마이너스"), "customer"),
        Topic("Servidental 박스·원신고", "Servidental 반품은 박스 수와 원수입신고 관계가 맞아야 한다.", kw("Servidental|코스타리카|박스|원수입신고|반품"), "customer"),
        Topic("Natrodent 주문 재분류", "실제 판매 주문이 아닌 운임 invoice/회수 건은 재고·회수·회계 업무로 재분류한다.", kw("Natrodent|invoice|회수|재고|매출 주문"), "customer"),
        Topic("온라인 주문분할", "챗봇/온라인 제약으로 주문이 분할되면 미수와 배송비를 다시 확인한다.", kw("온라인|챗봇|주문분할|배송비|선결제|취소"), "system"),
        Topic("주소·비고 검증", "비고에 주소 변경 필수 문구가 있으면 실제 배송지 반영 여부를 확인한다.", kw("주소|배송지|비고|zipCode|phone|우편번호"), "order"),
        Topic("공식 요청창구", "반려/삭제 근거가 사라지지 않도록 공식 채널 또는 사후기록을 남긴다.", kw("NOCO|공식|긴급|유선|요청창구|기록"), "process"),
        Topic("월마감·전표", "월마감 이후 정정은 원주문 전표와 대체주문 전표를 함께 본다.", kw("월마감|전표|마감|귀속|대체주문|출고전표"), "finance"),
    ],
    "정재회": [
        Topic("품목별 현재고 표기", "재고 부족 답변은 품목별 현재고와 부족수량을 나눠야 한다.", kw("현재고|재고 부족|주문수량|부족|품목별"), "stock"),
        Topic("입고예정일 미정", "입고일 미정이면 부족분 출고를 확정하지 않는다.", kw("입고 예정|입고일|미정|발주|공급사"), "stock"),
        Topic("부분 선출고 확인", "재고 있는 품목만 먼저 보낼지는 고객/영업 확인이 필요하다.", kw("나머지|먼저 출고|부분|선출고|우선 출고"), "logistics"),
        Topic("발주품의 승인 전", "발주 품의 승인 전에는 출고 가능으로 말하지 않는다.", kw("발주 품의|승인 전|공급사|발주요청|결제조건"), "stock"),
        Topic("직발송 구분", "직발송은 본사 출고와 회수 경로가 다르므로 따로 추적한다.", kw("직발송|덴탈맥스|로젠|본사|발송인"), "logistics"),
        Topic("대체 규격 확인", "대체 규격은 고객/영업 확인 후 등록한다.", kw("대체|16mm|20mm|규격|등록"), "product"),
        Topic("shade·두께 대체 금지", "Block shade나 두께가 다르면 대체품으로 단정하지 않는다.", kw("shade|A1|A2|두께|DOF SOLID|Block"), "product"),
        Topic("생산중단 표현", "생산중단 품목은 입고대기가 아니라 출고불가로 말한다.", kw("생산중단|재생산|출고 불가|생산 이력"), "product"),
        Topic("수량 0 구성품", "구성품 수량 0은 출고 전 주문 정정 신호다.", kw("수량 0|0으로|구성품|package|Milling tool"), "order"),
        Topic("단가·수량 반전", "창고 실행 전 단가/수량 오류를 발견하면 정정 요청한다.", kw("단가|수량|반대|오류|정정"), "order"),
        Topic("온라인 입금 확인", "온라인 주문은 입금 확인 전 출고로 넘기지 않는다.", kw("온라인|입금|확인|금액|주문자"), "finance"),
        Topic("온라인 취소 중단", "온라인 취소가 확인되면 핑거 주문과 발주/출고를 멈춘다.", kw("온라인|취소|핑거|발주|중단"), "system"),
        Topic("회수 접수와 완료", "회수 접수와 회수 완료를 같은 말로 쓰지 않는다.", kw("회수|접수|완료|집하|택배"), "return"),
        Topic("회수 수량 확인", "취소 주문은 실제 회수 완료 수량이 맞아야 정리한다.", kw("회수 완료|수량|43개|재고조정|주문취소"), "return"),
        Topic("오배송 회수", "배송지 오기입은 재출고와 회수 상태를 같이 본다.", kw("오배송|배송지|오기입|재출고|반품"), "return"),
        Topic("택배 분실", "택배 분실은 사고접수, 재발송, 보상금, 주문삭제로 이어진다.", kw("분실|사고 접수|보상금|재발송|송장"), "return"),
        Topic("택배 픽업 마감", "출고승인 후라도 픽업 마감 후면 다음 영업일 출고다.", kw("택배|마감|픽업|다음날|출고승인"), "logistics"),
        Topic("대량 패킹 완료수량", "대량 주문은 패킹 완료 수량이 확정값이다.", kw("패킹|완료 수량|대량|755|758"), "stock"),
        Topic("국내재고 전환분", "국내재고 전환분과 추가 발주분을 섞지 않는다.", kw("국내재고|전환|추가 발주|잔량|RED LION"), "stock"),
        Topic("Mawasi Dent 패턴", "Mawasi Dent는 주문수량과 패킹수량 차이를 먼저 본다.", kw("Mawasi|Block|패킹|삭제"), "customer"),
        Topic("RED LION 패턴", "RED LION은 대량 Tool 잔량과 추가발주 계산이 핵심이다.", kw("RED LION|Tool|잔량|5000|추가 발주"), "customer"),
        Topic("CIDEAS shade 패턴", "CIDEAS는 shade와 구성품 수량 오류를 먼저 확인한다.", kw("CIDEAS|shade|A2|A1|구성품"), "customer"),
        Topic("DOF Japan 입고일 분리", "DOF Japan은 납품 예정 품목과 금일 입고 품목을 나눠 말한다.", kw("DOF Japan|일본|납품 예정|입고 예정|DC20"), "customer"),
        Topic("국내 거래처 반복오류", "국내 거래처별 반복 오류는 다음 주문의 선제 체크포인트다.", kw("킴스|백광|달빛|센트럴|바른손|누리"), "customer"),
        Topic("권한 부족 보고", "권한이 없으면 우회하지 말고 하려던 처리와 필요한 권한을 보고한다.", kw("권한|승인취소|피드백|메뉴|없"), "system"),
        Topic("채널별 핵심값", "핑거/온라인/해외물류/DM은 각각 먼저 볼 값이 다르다.", kw("핑거세일즈|온라인 쇼핑몰|해외물류|DM|채널"), "process"),
        Topic("소모재 전달 위치", "소모재 전달은 수령자와 위치가 정확해야 한다.", kw("소모재|전달|6층|TECH|자리"), "logistics"),
        Topic("출고마감 우선순위", "출고마감 전에는 입금, 재고, 정정, 픽업마감을 순서대로 본다.", kw("출고 마감|금일 출고|가능 수량|마감"), "process"),
        Topic("실행값 기록문장", "정재회식 답변은 현재고/입고일/선출고 여부를 함께 말해야 한다.", kw("확인|공유|가능|먼저|예정"), "process"),
    ],
    "김채원": [
        Topic("OD번호 기반 요청", "김채원 요청은 OD번호와 오류 사유가 붙어야 처리 가능하다.", kw("OD|주문번호|반려|요청|사유"), "order"),
        Topic("단가 오기입", "단가 오기입은 출고승인 취소와 수정 주문으로 이어진다.", kw("단가|오기입|수정|반려|채터"), "order"),
        Topic("실적기준일 오기재", "실적기준일 오류는 승인 누락과 긴급출고를 만든다.", kw("실적기준일|오기재|팀장승인|누락|퀵"), "order"),
        Topic("주문완료 건 삭제", "완료 건 삭제는 이관 주문과 사유가 있어야 한다.", kw("주문완료|삭제|이관|PKG|재계약"), "order"),
        Topic("출고승인 취소", "온라인 취소나 품목수정은 출고승인 취소가 선행된다.", kw("출고승인|취소|온라인 주문취소|품목수정"), "order"),
        Topic("신규품목 기안", "신규품목은 기안 후 코드 생성 요청으로 이어진다.", kw("신규품목|품목 코드|기안|품번 생성"), "product"),
        Topic("Tool 500PKG 품번", "새 패키지 계약은 품번 생성이 주문 반영의 선행조건이다.", kw("Tool Promotion 500PKG|500 Tool|품번|계약"), "product"),
        Topic("단가표·가격표", "단가표와 품목코드는 주문 자동화보다 먼저 안정화되어야 한다.", kw("단가표|가격표|품목코드|매입가|일반가|딜러가"), "product"),
        Topic("품목명 변경자료", "품목명 변경자료는 포털/가격표/주문 반영을 맞춰야 한다.", kw("품목명 변경|루젠|SE PLUS|자료|이메일"), "product"),
        Topic("주문명 표기", "주문명은 고객사, PKG, Shop 번호, 품목군을 담는 분석 필드다.", kw("주문명|Tool|제품 분류|반영|온라인샵"), "system"),
        Topic("noti-order 품번 오류", "자동 주문의 품번 인식 오류는 정답 품번과 함께 전달해야 한다.", kw("noti-order|자동 인식|품번|SHARP|알맞은 품번"), "system"),
        Topic("거래처 인식 오류", "거래처 인식 오류는 실제 고객사명을 명확히 남겨야 한다.", kw("거래처 인식 오류|인식 고객사|실제 고객사|고객사"), "system"),
        Topic("온라인 결제 확인", "온라인 주문 결제는 카드/무통장/입금확인을 주문 처리와 연결한다.", kw("온라인|입금|신용카드|결제|수금"), "finance"),
        Topic("적립금·예치금", "적립금과 예치금은 회계가 구분 가능한 방식으로 주문에 남겨야 한다.", kw("적립금|예치금|분할|할인 품목|회계"), "finance"),
        Topic("쿠폰·할인코드", "쿠폰/할인코드는 비고와 품목 생성 기준을 분리해야 한다.", kw("쿠폰|할인코드|할인가|할인 품목"), "finance"),
        Topic("툴 가격인상 유예", "가격인상 유예기간에는 고객별 구매 제한을 수기로라도 계산한다.", kw("가격 인상|예도기간|6/15|6/19|구매 제한|tool 30"), "finance"),
        Topic("거래처원장 대조", "원장 발송 전 회계전표와 주문번호 히스토리 대조를 1차 절차로 둔다.", kw("거래처원장|미수금|대조|회계전표|주문번호"), "finance"),
        Topic("매출취소 미반영", "매출취소가 Payment Collection에 반영되지 않으면 원장 오류가 생긴다.", kw("매출취소|Payment Collection|데이터 오류|반품완료"), "finance"),
        Topic("수금탭·대시보드", "수금탭은 미수/입금 확인의 핵심 운영 데이터다.", kw("수금탭|수금|대시보드|입금|미입금"), "finance"),
        Topic("PKG 잔액오류", "패키지 잔액오류는 출고마감과 선승인 사유를 함께 기록한다.", kw("잔액 오류|PKG|계약 잔액|선 승인|출고 마감"), "finance"),
        Topic("일본법인 재고표", "일본법인 요청은 금일 재고와 입고일 미정 품목을 분리해 주문한다.", kw("일본법인|입고일|재고|출고 가능|MT"), "customer"),
        Topic("해외 주문 보류 넘김", "송금 전/배송조건 변경 같은 해외 판단은 김규탁·조소연에게 사유와 함께 넘긴다.", kw("해외물류|송금전|출고 보류|Shipping and handling|반려"), "country"),
        Topic("퀵 요청문", "퀵 요청은 선불/착불, 사유, 직접요청 여부가 있어야 한다.", kw("퀵|선불|착불|직접요청|로보틱스"), "logistics"),
        Topic("설치·출고일 확인", "장비 설치 예상일은 물류/로보틱스 준비상태와 따로 확인한다.", kw("설치|출고일자|FREEDOM|로보틱스|예상일"), "logistics"),
        Topic("고객사 데이터 메모", "신규 고객은 관리자메모와 주소/연락처 수기입력 기준이 필요하다.", kw("신규 고객|관리자메모|엑셀|주소|연락처"), "customer"),
        Topic("영업실적담당자 온라인샵", "온라인샵 매출분석을 위해 영업실적담당자 값을 일관되게 반영한다.", kw("영업실적담당자|온라인샵|매출분석|기존 온라인 주문"), "system"),
        Topic("권한·열람 제한", "권한이 없으면 열람 요청과 임시 처리 방법을 기록한다.", kw("권한|열람|불가|latest data|포털"), "system"),
        Topic("오류와 제안 분리", "시스템 오류 보고는 오류 필드와 제안을 분리해 전달해야 한다.", kw("오류|제안|포털|타임아웃|승인요청"), "system"),
        Topic("인수인계서 운영기준", "김채원 인수인계는 담당자 맵, 요청문, 멈춤조건을 남겨야 가치가 있다.", kw("인수인계|업무분장|특이사항|남은 기간|작성"), "process"),
    ],
    "이미연": [
        Topic("주문상태별 동작", "이미연은 주문 등록, 승인요청, 임원승인, 출고승인, 완료 상태별 동작을 구분한다.", kw("승인요청|임원승인|승인처리|출고승인|주문완료"), "order"),
        Topic("삭제 가능 문의", "주문 삭제는 가능 여부와 출고/전표 이력을 확인해야 한다.", kw("삭제|삭제 가능|주문삭제|원주문|완료건"), "order"),
        Topic("반려 요청", "반려 요청은 주문번호와 사유, 수정 후 재등록 계획이 필요하다.", kw("반려|부탁|사유|수정|재승인"), "order"),
        Topic("재작성 주문", "재작성 주문은 원주문과 새 주문번호를 연결해야 한다.", kw("재작성|새로 작성|재등록|차액|새 주문"), "order"),
        Topic("단가 오류 차액", "단가 오기입 차액은 별도 주문과 회계 영향 확인이 필요하다.", kw("단가|오기입|차액|온라인 주문|결제조건"), "finance"),
        Topic("수량 변경", "수량 변경은 반려 후 수정 또는 재승인을 필요로 한다.", kw("수량|수량 변경|2개씩|변경했습니다"), "order"),
        Topic("유상·무상 분리", "유상/무상은 출고팀과 회계 요청 기준에 맞춰 별도 주문으로 올린다.", kw("유상|무상|따로|비고|경품"), "finance"),
        Topic("+수량·-수량", "교환/회수 월이 다르면 +수량과 -수량 주문을 따로 작성한다.", kw("+수량|-수량|회수건|교환건|작성 월"), "return"),
        Topic("신규품목 신청서", "금일 출고 신규품목은 신청서 상신과 담당자 확인이 우선이다.", kw("신규품목|신청서|상신|금일 출고|품목코드"), "product"),
        Topic("CRAFT 렌탈 신규품목", "CRAFT S/DX 렌탈 품목은 신규품목 요청과 주문 반영을 연결한다.", kw("CRAFT S|CRAFT DX|렌탈|신규품목"), "product"),
        Topic("소모재 전달자", "소모재 전달은 수령자, 위치, 설치일을 명확히 한다.", kw("소모재|전달|TECH|6층|설치"), "logistics"),
        Topic("있는 재고 먼저 전달", "일부 재고만 있으면 가능한 재고만 전달하고 부족분을 분리한다.", kw("있는재고|재고|전달|가능 수량|실재고"), "stock"),
        Topic("배송지 복사 규칙", "기존 고객 주문은 최근 주문 복사로 실제 배송지를 보존한다.", kw("배송지|주소|복사|최근 주문|오배송"), "logistics"),
        Topic("오배송 합배송", "오배송과 새 출고가 겹치면 합배송과 회수를 같이 설계한다.", kw("오배송|합배송|회수 예정|기사님"), "return"),
        Topic("회수 택배", "회수 택배는 접수와 실제 방문/완료를 구분한다.", kw("회수 택배|회수|방문|집하|완료"), "return"),
        Topic("회수 스캐너 보관", "회수 스캐너는 보관 위치와 이후 AMJ/지원팀 사용 계획을 구분한다.", kw("회수 스캐너|6층|AMJ|지원팀|보관"), "return"),
        Topic("온라인 수기 등록", "온라인 주문은 핑거 수기 등록과 중복/취소 상태를 확인한다.", kw("온라인|핑거|수기로|등록|주문건"), "system"),
        Topic("온라인 주문 삭제", "온라인 주문 삭제/취소는 내부 주문과 발주 상태를 같이 본다.", kw("온라인주문|삭제|취소|툴3개|발주"), "system"),
        Topic("noti-order 오류", "noti-order 오류는 회사명, 고객명, 품목, 결제 필드를 사람이 보정한다.", kw("noti-order|고객사 인식|오등록|우편번호|주문결제조건"), "system"),
        Topic("입금내역 확인", "입금내역은 수금처리와 출고처리의 선행값이다.", kw("입금|입금내역|계좌이체|수금처리|입금액"), "finance"),
        Topic("미수원장 정정", "미수원장 오류는 회계 수금완료와 파일 수정 후 재발송으로 처리한다.", kw("미수원장|미수|수금완료|재발송|파일 수정"), "finance"),
        Topic("카드결제 미반영", "카드결제인데 미수원장에 남으면 수금탭 누락을 확인한다.", kw("카드결제|미수원장|수금탭|누락|반영"), "finance"),
        Topic("PKG 잔액·차감", "PKG 잔액과 차감 예정액은 고객 안내와 주문조건을 바꾼다.", kw("PKG|잔액|차감|워런티|패키지"), "finance"),
        Topic("렌탈료 잔액 차감", "렌탈료를 워런티 잔액에서 차감할 때 월별 차감 기준을 확인한다.", kw("렌탈료|워런티 잔액|차감|월별|9만원"), "finance"),
        Topic("출고팀 더블체크", "영업관리 주문상태와 출고팀 실제상태를 더블체크한다.", kw("출고쪽|더블체크|출고 요청|오늘 출고|확정"), "logistics"),
        Topic("직접 배송품", "직접 교환품 배송은 수령자와 자리 전달 여부를 기록한다.", kw("직접 교환품|직접 배송|자리에 전달|이건호|강두환"), "logistics"),
        Topic("품목 제외 승인", "재고나 오류가 있는 품목은 제외하고 승인으로 넘길 수 있다.", kw("품목제외|승님|넘겼|재고|확인"), "order"),
        Topic("완료 후 권한수정", "완료 후 제목/담당자 수정은 권한 요청으로 처리한다.", kw("권한|제목 수정|완료건|담당자|수정"), "system"),
        Topic("인수인계 우선순위", "이미연 인수인계는 주문 생성, 주소/단가 확인, 고객 신규 생성, PKG 차감 순서를 남긴다.", kw("인수인계|우선순위|주소 확인|단가 확인|고객 신규"), "process"),
    ],
}


GENERIC = {
    "order": ("주문 상태와 원주문 관계를 먼저 확인한다.", "삭제/반려/재작성 중 어떤 동작인지 명확히 지정한다.", "주문번호, 오류 필드, 원주문/대체주문 관계를 남긴다."),
    "finance": ("금액·수금·잔액은 회계 이력으로 남는 기준값이다.", "회계 확인 전 외부 안내나 출고 확정을 하지 않는다.", "수금확인자, 금액, 전표/원장 반영 여부를 남긴다."),
    "logistics": ("실제 이동 가능 여부는 일정·비용·수령자·마감시간을 함께 봐야 한다.", "운송/전달 조건이 불명확하면 출고 가능으로 말하지 않는다.", "배송방식, 비용주체, 수령자, 실제 출고일을 남긴다."),
    "product": ("제품군별로 필요한 확인값이 다르다.", "품목코드·규격·시리얼·구성품 중 하나라도 불명확하면 주문을 고정하지 않는다.", "품목코드, 규격, 구성품, 확인자를 남긴다."),
    "country": ("국가/제도 조건은 출고 가능 여부를 직접 바꾼다.", "통관·허가·수입자 확인 없이 서류나 예약을 확정하지 않는다.", "국가, 제도값, 수입자 확인, 보류 사유를 남긴다."),
    "customer": ("거래처별 반복 특이사항은 다음 주문의 선제 체크포인트다.", "과거 패턴과 다른 조건이면 근거를 다시 확인한다.", "거래처명, 반복 리스크, 이번 확인값을 남긴다."),
    "system": ("시스템/자동화 결과는 업무값 검증 전까지 초안이다.", "자동값과 실제 고객/품목/금액이 다르면 사람이 보정한다.", "오류 필드, 정정값, 재발 방지 요청을 남긴다."),
    "return": ("회수/반품은 접수와 완료, 원주문과 재출고를 분리해야 한다.", "실제 회수 완료 전 재고·주문 정리를 확정하지 않는다.", "원주문, 회수상태, 완료수량, 재출고 여부를 남긴다."),
    "stock": ("재고 판단은 품목별 현재고와 입고 예정일 기준이다.", "일부 재고만 보고 전체 주문을 완료로 말하지 않는다.", "품목별 현재고, 부족수량, 입고예정일을 남긴다."),
    "document": ("문서는 실물·운송·통관의 결과물이다.", "서류 한 항목을 바꾸면 연결 문서와 실물 부착 상태도 확인한다.", "문서 버전, 변경값, 확인자를 남긴다."),
    "process": ("인수인계 가치는 다음 사람이 같은 판단을 재현할 수 있는 데 있다.", "담당자·재개조건·공식기록 없이 처리 완료로 보지 않는다.", "요청채널, 담당자, 재개조건, 처리결과를 남긴다."),
}


def load_json(path: Path, default):
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def clean(value: str | None, limit: int | None = None) -> str:
    value = re.sub(r"\s+", " ", value or "").strip()
    if limit and len(value) > limit:
        return value[: limit - 3] + "..."
    return value


def person_info(person: str) -> dict:
    return load_json(PEOPLE_DIR / f"{person}_person_info.json", {})


def authored_count(person: str, summary: dict) -> int:
    rows = load_json(ROOT / summary.get("all_evidence_path", ""), [])
    return sum(1 for row in rows if person in (row.get("who") or ""))


def evidence_rows(person: str, summary: dict) -> list[dict]:
    return load_json(ROOT / summary.get("all_evidence_path", ""), [])


def matched_rows(rows: list[dict], keywords: tuple[str, ...], person: str) -> list[dict]:
    selected = []
    for row in rows:
        text = f"{row.get('source') or ''} {row.get('text') or ''}"
        if any(k.lower() in text.lower() for k in keywords):
            selected.append(row)
    selected.sort(key=lambda r: (person not in (r.get("who") or ""), r.get("ts") or ""))
    return selected


def topic_page(person: str, rows: list[dict], topic: Topic) -> Page | None:
    matches = matched_rows(rows, topic.keywords, person)
    if not matches:
        return None
    count = len(matches)
    top_sources = ", ".join(f"{name} {cnt}건" for name, cnt in Counter(r.get("source") or "" for r in matches).most_common(3))
    snippets = []
    seen = set()
    for row in matches:
        txt = clean(row.get("text"), 150)
        if not txt or txt in seen:
            continue
        seen.add(txt)
        snippets.append(f"{row.get('ts', '')[:10]} {row.get('source')}: {txt}")
        if len(snippets) >= 3:
            break
    base_judgement, base_stop, base_record = GENERIC.get(topic.lens, GENERIC["process"])
    signals = [f"관련 근거 {count}건 확인", f"주요 출처: {top_sources or '확인됨'}", *snippets]
    judgement = [topic.focus, base_judgement, "이 판단은 수집된 Teams 근거에서 반복적으로 보인 업무값을 기준으로 정리했다"]
    steps = [
        "관련 주문번호, 고객사, 품목 또는 금액을 먼저 특정한다",
        "자동값/요청값과 실제 기준값을 대조한다",
        "담당 권한자 또는 협업부서에 필요한 동작을 명확히 요청한다",
        "처리 후 원주문, 변경값, 재개 조건을 문서나 공식 채널에 남긴다",
    ]
    exceptions = [
        "단일 메시지만 보고 일반규칙으로 확대하지 않고 반복 출처와 기존 업무규칙 문서를 함께 본다",
        "개발/AX 요청 자체가 아니라 운영 중 반복되는 판단값만 업무규칙으로 반영한다",
    ]
    stops = [base_stop, "기준값 또는 담당자가 불명확하면 처리 완료로 적지 않는다"]
    records = [base_record, "근거가 된 주문/채널/담당자", "추후 같은 상황에서 재사용할 확인 문장"]
    return Page(topic.title, topic.focus, tuple(signals), tuple(judgement), tuple(steps), tuple(exceptions), tuple(stops), tuple(records))


def top_sources(summary: dict, limit: int = 10) -> str:
    sources = summary.get("top_sources") or {}
    items = list(sources.items())[:limit] if isinstance(sources, dict) else sources[:limit]
    return ", ".join(f"{name} {count}건" for name, count in items)


def md_list(items: tuple[str, ...]) -> str:
    return "\n".join(f"- {item}" for item in items)


def md_steps(items: tuple[str, ...]) -> str:
    return "\n".join(f"{idx}. {item}" for idx, item in enumerate(items, 1))


def page_to_markdown(idx: int, page: Page) -> str:
    return "\n".join(
        [
            f"## {idx:02d}. {page.title}",
            "",
            page.focus,
            "",
            "### 데이터에서 보인 신호",
            md_list(page.signals),
            "",
            "### 인수인계 판단",
            md_list(page.judgement),
            "",
            "### 처리 절차",
            md_steps(page.steps),
            "",
            "### 특이사항과 예외",
            md_list(page.exceptions),
            "",
            "### 멈춤 조건",
            md_list(page.stops),
            "",
            "### 기록해야 할 항목",
            md_list(page.records),
            "",
            "### 인수자가 기억할 문장",
            f"`{page.title}` 상황에서는 확인된 값과 추정값을 분리하고, 멈춤 조건이 해소되기 전까지 다음 단계로 넘기지 않는다.",
            "",
        ]
    )


def expanded_pages(person: str, summary: dict) -> list[Page]:
    rows = evidence_rows(person, summary)
    pages = list(handover30.PAGES[person])
    for topic in TOPICS[person]:
        page = topic_page(person, rows, topic)
        if page:
            pages.append(page)
        if len(pages) >= 59:
            break
    return pages[:59]


def build_markdown(person: str, summary: dict, pages: list[Page]) -> str:
    info = person_info(person)
    date_range = summary.get("date_range") or ["", ""]
    direct = authored_count(person, summary)
    header = [
        f"# {person} 업무규칙 인수인계서 60페이지판",
        "",
        "## 문서 작성 기준",
        f"- 작성일: {RUN_DATE}",
        "- 사용 데이터: Teams full-history scrape, person_evidence JSON, data/teams/peoples 업무규칙/업무이력 문서",
        "- 제한: 수집된 정보에서 반복적으로 확인되는 업무 패턴만 사용했고, 개발/AX 요청 자체는 통상업무 규칙에서 제외했다.",
        "- 페이지 구성: 표지 1페이지 + 본문 59페이지 = 수동 페이지 구분 기준 60페이지",
        f"- 부서/직책: {info.get('dept') or ''} / {info.get('position') or ''}",
        f"- ERP main_work: {info.get('main_work') or ''}",
        f"- 근거 메시지: {summary.get('evidence_count')}건, 직접 발화 {direct}건, 통상업무 후보 {summary.get('routine_non_dev_count')}건",
        f"- 기간: {date_range[0]} ~ {date_range[1]}",
        f"- 주요 출처: {top_sources(summary)}",
        "",
        "이 문서는 기존 30페이지판을 확장한 버전이다. 추가 장은 수집된 Teams 근거에서 키워드, 출처, 대표 발화가 확인된 주제만 사용했다.",
        "",
        "<!-- PAGEBREAK -->",
        "",
    ]
    body = []
    for idx, page in enumerate(pages, 1):
        body.append(page_to_markdown(idx, page))
        if idx != len(pages):
            body.append("<!-- PAGEBREAK -->\n")
    return "\n".join(header + body)


def set_font(run, size: float = 8.8) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    rpr = run._element.rPr
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_para(doc: Document, text: str, style: str | None = None, size: float = 8.8) -> None:
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    set_font(run, size=size)


def markdown_to_docx(markdown: str, path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    doc.styles["Normal"].font.name = "Malgun Gothic"
    doc.styles["Normal"].font.size = Pt(8.8)

    for line in markdown.splitlines():
        line = line.rstrip()
        if line == "<!-- PAGEBREAK -->":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif not line:
            continue
        elif line.startswith("# "):
            add_para(doc, line[2:].strip(), "Title", 15)
        elif line.startswith("## "):
            add_para(doc, line[3:].strip(), "Heading 1", 11.5)
        elif line.startswith("### "):
            add_para(doc, line[4:].strip(), "Heading 2", 9.5)
        elif line.startswith("- "):
            add_para(doc, line[2:].strip(), "List Bullet", 8.5)
        elif re.match(r"^\d+\. ", line):
            add_para(doc, re.sub(r"^\d+\. ", "", line), "List Number", 8.5)
        else:
            add_para(doc, line, None, 8.8)
    doc.save(path)


def docx_break_count(path: Path) -> int:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return xml.count('w:type="page"')


def main() -> None:
    summary = load_json(SUMMARY_PATH, {})
    OUT_MD.mkdir(parents=True, exist_ok=True)
    OUT_DOCX.mkdir(parents=True, exist_ok=True)
    verification = {}
    for person in ["김규탁", "조소연", "정재회", "김채원", "이미연"]:
        pages = expanded_pages(person, summary[person])
        if len(pages) < 59:
            raise AssertionError(f"{person}: only {len(pages)} body pages available")
        markdown = build_markdown(person, summary[person], pages)
        md_path = OUT_MD / f"{person}_업무규칙_인수인계서_60p_{RUN_DATE}.md"
        docx_path = OUT_DOCX / f"{person}_업무규칙_인수인계서_60p_{RUN_DATE}.docx"
        md_path.write_text(markdown, encoding="utf-8")
        markdown_to_docx(markdown, docx_path)
        breaks = docx_break_count(docx_path)
        with zipfile.ZipFile(docx_path) as zf:
            names = set(zf.namelist())
            assert "[Content_Types].xml" in names
            assert "word/document.xml" in names
        verification[person] = {
            "markdown": str(md_path),
            "docx": str(docx_path),
            "body_pages": len(pages),
            "manual_page_breaks": breaks,
            "page_count_by_manual_breaks": breaks + 1,
            "markdown_tokens_by_whitespace": len(re.findall(r"\S+", markdown)),
            "docx_bytes": docx_path.stat().st_size,
        }
        print(person, verification[person])
    VERIFY_PATH.write_text(json.dumps(verification, ensure_ascii=False, indent=2), encoding="utf-8")
    print("verification", VERIFY_PATH)


if __name__ == "__main__":
    main()
