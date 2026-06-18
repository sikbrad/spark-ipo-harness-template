#!/usr/bin/env python3
"""Build one DOCX per person for Teams business-rule analysis."""
from __future__ import annotations

import importlib.util
import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


RUN_DATE = "2026-06-17"
ROOT = Path("data/teams/full-history") / RUN_DATE
OUT = Path("output/teams-business-rules") / RUN_DATE / "individual_docs"
PEOPLE_DIR = Path("data/company/people")
SUMMARY_PATH = Path("output/teams-business-rules") / RUN_DATE / "evidence_summary.json"
BUILD_SCRIPT = Path("proc/plan/build_teams_business_rules_docx_2026_06_17.py")

spec = importlib.util.spec_from_file_location("combined_report", BUILD_SCRIPT)
combined = importlib.util.module_from_spec(spec)
assert spec.loader
spec.loader.exec_module(combined)

TARGETS = ["김규탁", "조소연", "정재회", "김채원", "이미연"]

WORKFLOWS = {
    "김규탁": [
        {
            "request": "해외 출고/출하 요청",
            "intake": "주문번호, 국가/업체명, 품목, 출하조건, 수금/승인 상태, 요청일",
            "checks": "주문 내용과 품목/수량, 결재 단계, 미수/후불 리스크, CIPL/인보이스/운임/통관 조건",
            "process": "근거가 맞으면 출고 일정과 서류를 진행하고, 근거가 부족하면 보류/반려 후 담당자에게 재요청하게 한다.",
            "output": "출고 승인/보류/반려 기록, 서류 업데이트, 담당자에게 확인 요청",
        },
        {
            "request": "미수·후불·예외 출고",
            "intake": "미수 상태, 수금 예정/증빙, 본부장/대표 승인 기록, 업무기안",
            "checks": "구두 승인인지, 결재방/기안/대표 승인처럼 추적 가능한 근거가 있는지",
            "process": "구두 승인만 있으면 출고 근거로 인정하지 않고, 영업부가 공식 승인 근거를 남긴 뒤 처리한다.",
            "output": "출고보류 또는 승인 근거 확인 후 출고",
        },
        {
            "request": "품목/단가/수량 오류 또는 오발송",
            "intake": "OD/Sales 번호, 고객명, 오류 품목, 실제 출고/회수 상태",
            "checks": "주문서와 실제 품목, 출고 이력, 회수 필요 여부",
            "process": "오류를 주문번호 기준으로 공유하고 반려/수정/회수/부분 출고를 결정한다.",
            "output": "반려 요청, 수정 재등록 요청, 회수 택배/재발송 안내",
        },
        {
            "request": "해외영업-물류 소통 정리",
            "intake": "채팅/메신저/유선으로 흩어진 요청",
            "checks": "공식 요청창구에 남아 있는지, 긴급 건 사후 기록이 있는지",
            "process": "NOCO 등 공식 창구로 모으고 단순 메신저 요청을 업무 기록으로 전환한다.",
            "output": "요청/처리 이력 누락 감소, 출고 판단 근거 축적",
        },
    ],
    "조소연": [
        {
            "request": "주문 등록/수정/반려/삭제",
            "intake": "주문번호, 고객/국가, 품목, 수정 사유, 요청자",
            "checks": "주문 상태, 반려/삭제 가능 여부, 회계마감일/출고일/운송장 같은 이력 필드",
            "process": "요청 사유를 기준으로 반려·삭제·신규 재작성 여부를 결정하고 이력이 남게 한다.",
            "output": "반려/삭제/재작성 안내, 이력 추적 기준",
        },
        {
            "request": "출하 근거 확인",
            "intake": "수금등록, 송금증빙, 주문번호, 출하 요청 사유",
            "checks": "회계팀 수금 등록 여부와 송금증빙, 보류 사유에 대한 공식 근거",
            "process": "근거가 없으면 출하 요청을 보류하고, 회계 확인과 증빙 보강 후 재요청하게 한다.",
            "output": "출하 승인 또는 근거 부족 반려",
        },
        {
            "request": "출고 지연/재고 부족/직발송",
            "intake": "발송요청일, 발송예정일, 지연 사유, 매입처 재고, 직발송 가능 여부",
            "checks": "매입처 재고와 당사 입고 필요 여부, 부분 출고 가능 여부",
            "process": "지연 사유를 주문번호 기준으로 공지하고 담당자에게 일정/부분 출고 판단을 받는다.",
            "output": "지연 안내, 일정 업데이트, 주문 재작성 또는 부분 출고",
        },
        {
            "request": "공식 요청창구/긴급건 운영",
            "intake": "NOCO, 채팅방, 유선, 메신저로 들어온 요청",
            "checks": "요청사항과 해결이 공식 기록에 남았는지, 긴급 건은 유선 더블체크가 되었는지",
            "process": "일반 요청은 공식 창구로 옮기고, 긴급 건은 유선 확인 후 사후 기록을 남긴다.",
            "output": "요청 누락 방지, 긴급 처리 이력 확보",
        },
    ],
    "정재회": [
        {
            "request": "재고 부족/부분 출고 확인",
            "intake": "주문번호, 고객명, 품목, 부족 수량, 가능한 입고/출고일",
            "checks": "현재 재고, 입고 예정, 나머지 품목 선출고 가능 여부",
            "process": "부족 품목과 가능한 출고일을 공유하고 담당자에게 부분 출고 여부를 확인한다.",
            "output": "부분 출고, 전체 보류, 출고 예정일 안내",
        },
        {
            "request": "소모재/장비 전달",
            "intake": "주문명, 수령자, 전달 장소, 설치일/필요일",
            "checks": "전달 대상자가 누구인지, 자리 부재 시 보관 위치가 명확한지",
            "process": "수령자 또는 팀에 전달하고, 부재 시 어디에 두었는지 기록한다.",
            "output": "전달 완료 기록, 추가 확인 사항",
        },
        {
            "request": "회수/재발송/오출고",
            "intake": "오출고 내용, 회수 접수 상태, 송장번호, 재발송 필요 여부",
            "checks": "회수 접수가 실제로 되었는지, 택배사 연락이 필요한지, 새 송장번호가 있는지",
            "process": "회수 접수/취소/재발송 상태를 정리하고 송장번호를 공유한다.",
            "output": "회수/재발송 완료, 송장번호 공유",
        },
        {
            "request": "반려/승인취소 권한 문제",
            "intake": "처리하려는 주문 상태와 필요한 권한",
            "checks": "본인 계정으로 반려/승인취소가 가능한지",
            "process": "권한이 없으면 즉시 시스템 권한 부여 요청을 올린다.",
            "output": "권한 보완 요청, 처리 지연 원인 공유",
        },
    ],
    "김채원": [
        {
            "request": "국내 주문 오류 정정",
            "intake": "주문번호, 주문명, 단가/유무상/담당자/고객 정보 오류",
            "checks": "기존 주문의 오류 위치, 반려 필요 여부, 수정 주문 작성 가능 여부",
            "process": "오류 사유를 명시해 반려를 요청하고 수정 주문을 재등록한다.",
            "output": "반려 요청, 수정 주문, 담당자 확인",
        },
        {
            "request": "신규 품목/품목 코드 생성",
            "intake": "신규 품목명, 기안 상태, 품목 코드 생성 필요 사유",
            "checks": "기안이 올라갔는지, 주문 입력 전에 코드가 필요한지",
            "process": "기안 이후 품목 코드 생성 담당자에게 요청하고 생성 후 주문에 반영한다.",
            "output": "품목 코드 생성 요청, 주문 반영",
        },
        {
            "request": "빠른 출고/퀵 발송",
            "intake": "주문번호, 주문명, 품목, 선불/착불, 급한 사유",
            "checks": "출고 가능 여부, 담당팀 직접 요청 여부, 비용 부담 방식",
            "process": "물류/출고 담당자에게 주문번호와 사유를 붙여 명확히 넘긴다.",
            "output": "퀵 발송 요청, 출고 진행 확인",
        },
        {
            "request": "온라인샵 정책 운영",
            "intake": "적립금/예치금, 할인코드, 가격 인상 유예, 구매 제한",
            "checks": "정책이 실제 주문 입력 기준으로 어떻게 바뀌는지",
            "process": "정책을 주문 처리 규칙으로 번역하고 운영/자동화 쪽에 반영 요청한다.",
            "output": "온라인샵 주문 처리 기준, 예외 처리 기준",
        },
    ],
    "이미연": [
        {
            "request": "국내 주문 등록/재작성/반려",
            "intake": "주문번호, 주문 오류, 신규품목 여부, 승인 단계",
            "checks": "삭제/반려 가능 여부, 임원승인 단계, 수정 후 재등록 필요 여부",
            "process": "먼저 수정·재등록하고, 필요한 경우 조소연/김규탁에게 삭제·반려·승인 확인을 요청한다.",
            "output": "수정 주문, 반려 요청, 승인 단계 확인",
        },
        {
            "request": "소모재/장비 전달 요청",
            "intake": "고객명, 품목, 전달 대상자, 설치일, 재고 여부",
            "checks": "물류 담당자가 전달할 수 있는지, 수령자가 누구인지",
            "process": "정재회 등 물류 담당자에게 전달 대상과 장소를 구체적으로 안내한다.",
            "output": "소모재 전달 요청, 수령자 안내",
        },
        {
            "request": "온라인샵 주문 보정",
            "intake": "noti-order 알림, 회사명/고객명/주문명/품목/결제정책 오류",
            "checks": "자동화 입력값과 실제 고객/주문 정보가 맞는지",
            "process": "자동화 결과를 사람이 최종 보정하고 필요한 권한/수정 요청을 별도로 올린다.",
            "output": "보정된 주문, 권한 요청, 오류 공유",
        },
        {
            "request": "업무 인수인계",
            "intake": "주문 처리 파일, 예외처리, 특이사항, 순서",
            "checks": "신규 담당자가 어떤 파일을 보고 어떤 순서로 처리해야 하는지",
            "process": "파일 위치, 일반 처리 순서, 예외 A/B 처리 방식을 문서화한다.",
            "output": "인수인계 문서, 업무 처리 기준",
        },
    ],
}


def load_json(path: Path, default=None):
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def person_info(person: str) -> dict:
    return load_json(PEOPLE_DIR / f"{person}_person_info.json", {})


def clean(value: str, limit: int | None = None) -> str:
    value = re.sub(r"\s+", " ", value or "").strip()
    if limit and len(value) > limit:
        return value[: limit - 3] + "..."
    return value


def person_authored_samples(summary: dict, person: str, limit: int = 24) -> list[dict]:
    path = Path(summary[person]["all_evidence_path"])
    rows = load_json(path, [])
    selected = []
    seen = set()
    for row in rows:
        if "author" not in (row.get("reasons") or []):
            continue
        if not row.get("routine_categories"):
            continue
        if row.get("is_dev_system") and len(row.get("routine_categories") or []) < 2:
            continue
        key = (row.get("source"), row.get("text", "")[:80])
        if key in seen:
            continue
        seen.add(key)
        selected.append(row)
    selected.sort(key=lambda r: r.get("ts") or "")

    # Keep a spread across categories/sources rather than only oldest messages.
    out = []
    used_categories = set()
    for row in selected:
        cats = row.get("routine_categories") or []
        if any(cat not in used_categories for cat in cats):
            out.append(row)
            used_categories.update(cats)
        if len(out) >= limit // 2:
            break
    for row in selected[-limit:]:
        if row not in out:
            out.append(row)
        if len(out) >= limit:
            break
    return out[:limit]


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def render_markdown(person: str, summary: dict, analysis: dict, info: dict, samples: list[dict]) -> str:
    stat = summary[person]
    lines = [
        f"# {person} 업무규칙 분석",
        "",
        f"- 작성일: {RUN_DATE}",
        f"- 원천 데이터: `{ROOT}`",
        f"- 부서/직책: {info.get('dept', '')} / {info.get('position', '')}",
        f"- ERP main_work: {info.get('main_work')}",
        f"- 근거 메시지: {stat.get('evidence_count'):,}건",
        f"- 통상업무 후보: {stat.get('routine_non_dev_count'):,}건",
        f"- 기간: {stat.get('date_range', ['',''])[0]} ~ {stat.get('date_range', ['',''])[1]}",
        "",
        "## 역할 요약",
        "",
        analysis["role_summary"],
        "",
        "## 실제 수행 업무",
        "",
    ]
    lines.extend(f"- {item}" for item in analysis["work_done"])
    lines.extend(["", "## 요청 유형별 처리 방식", ""])
    for flow in WORKFLOWS[person]:
        lines.extend([
            f"### {flow['request']}",
            "",
            f"- 접수 정보: {flow['intake']}",
            f"- 확인 기준: {flow['checks']}",
            f"- 처리 방식: {flow['process']}",
            f"- 산출/기록: {flow['output']}",
            "",
        ])
    lines.extend(["## 업무규칙", ""])
    lines.extend(f"- {item}" for item in analysis["rules"])
    lines.extend(["", "## 근거 예시", ""])
    lines.extend(f"- {item}" for item in analysis["evidence"])
    lines.extend(["", "## 작성자 기준 메시지 샘플", ""])
    for row in samples:
        lines.append(
            f"- {row.get('ts')} | {row.get('source')} | "
            f"{'/'.join(row.get('routine_categories') or [])}: {clean(row.get('text') or '', 260)}"
        )
    return "\n".join(lines) + "\n"


def render_docx(person: str, summary: dict, analysis: dict, info: dict, samples: list[dict]) -> Path:
    stat = summary[person]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(9.5)

    doc.add_heading(f"{person} 업무규칙 분석", 0)
    add_table(doc, ["항목", "내용"], [
        ["작성일", RUN_DATE],
        ["부서/직책", f"{info.get('dept', '')} / {info.get('position', '')}"],
        ["ERP main_work", str(info.get("main_work"))],
        ["근거 메시지", f"{stat.get('evidence_count'):,}건"],
        ["통상업무 후보", f"{stat.get('routine_non_dev_count'):,}건"],
        ["분석 기간", f"{stat.get('date_range', ['',''])[0]} ~ {stat.get('date_range', ['',''])[1]}"],
        ["주요 출처", ", ".join(f"{k}({v})" for k, v in list(stat.get("top_sources", {}).items())[:8])],
        ["상위 카테고리", ", ".join(f"{k}({v})" for k, v in list(stat.get("category_counts", {}).items())[:8])],
    ])

    doc.add_heading("1. 역할 요약", 1)
    doc.add_paragraph(analysis["role_summary"])

    doc.add_heading("2. 실제 수행 업무", 1)
    add_bullets(doc, analysis["work_done"])

    doc.add_heading("3. 요청 유형별 처리 방식", 1)
    add_table(
        doc,
        ["요청 유형", "접수 정보", "확인 기준", "처리 방식", "산출/기록"],
        [[flow["request"], flow["intake"], flow["checks"], flow["process"], flow["output"]] for flow in WORKFLOWS[person]],
    )

    doc.add_heading("4. 업무규칙", 1)
    add_bullets(doc, analysis["rules"])

    doc.add_heading("5. 근거 예시", 1)
    add_bullets(doc, analysis["evidence"])

    doc.add_heading("6. 작성자 기준 메시지 샘플", 1)
    add_table(
        doc,
        ["시각", "출처", "카테고리", "내용"],
        [
            [
                row.get("ts") or "",
                row.get("source") or "",
                ", ".join(row.get("routine_categories") or []),
                clean(row.get("text") or "", 360),
            ]
            for row in samples
        ],
    )

    doc.add_heading("7. 제외 기준", 1)
    add_bullets(doc, [
        "개발 요청, 포털/CRM/AX 기능 구현 요청, 버그 수정 요청은 통상업무 규칙에서 제외했다.",
        "다만 시스템 요구가 운영 절차를 설명하는 근거로 쓰인 경우에는 처리방식의 배경으로만 반영했다.",
        "메시지 후보는 Teams 전체 수집 raw에서 사람 작성/언급/참여방 근거를 모은 뒤 통상업무 카테고리 중심으로 좁혔다.",
    ])

    path = OUT / f"{person}_업무규칙_분석_{RUN_DATE}.docx"
    doc.save(path)
    return path


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    summary = load_json(SUMMARY_PATH, {})
    result = []
    for person in TARGETS:
        analysis = combined.PERSON_ANALYSIS[person]
        info = person_info(person)
        samples = person_authored_samples(summary, person)
        md = render_markdown(person, summary, analysis, info, samples)
        md_path = OUT / f"{person}_업무규칙_분석_{RUN_DATE}.md"
        md_path.write_text(md, encoding="utf-8")
        docx_path = render_docx(person, summary, analysis, info, samples)
        result.append({"person": person, "docx": str(docx_path), "markdown": str(md_path)})
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
