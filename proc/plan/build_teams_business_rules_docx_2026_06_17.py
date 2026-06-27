#!/usr/bin/env python3
"""Build the Teams full-history/business-rules report as Markdown and DOCX."""
from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt


RUN_DATE = "2026-06-17"
ROOT = Path("data/teams/full-history") / RUN_DATE
OUT = Path("output/teams-business-rules") / RUN_DATE
REPORT_MD = OUT / f"teams_full_history_business_rules_{RUN_DATE}.md"
REPORT_DOCX = OUT / f"teams_full_history_business_rules_{RUN_DATE}.docx"
PEOPLE_DIR = Path("data/company/people")

TARGETS = ["김규탁", "조소연", "정재회", "김채원", "이미연"]

CATEGORY_KEYWORDS = {
    "주문/매출": ["주문", "매출", "견적", "승인", "핑거", "세일즈", "영업", "쇼핑몰", "OD20", "Sales"],
    "출고/물류": ["출고", "출하", "배송", "송장", "택배", "물류", "CIPL", "인보이스", "통관", "재고", "입고", "픽업", "납기", "포장", "창고"],
    "정산/회계": ["입금", "결제", "세금", "계산서", "정산", "미수", "청구", "마감", "회계", "CMS", "카드", "환불"],
    "수출/통관": ["수출", "수입", "관세", "HS", "전략물자", "원산지", "DAP", "EXW", "DHL", "FedEx", "UPS"],
    "요청/확인": ["요청", "확인", "부탁", "전달", "공유", "문의", "처리", "안내", "가능", "반려", "보류"],
    "개발/시스템": ["개발", "API", "서버", "버그", "배포", "CRM", "AX", "포탈", "앱", "자동화", "오류", "시스템"],
}

PERSON_ANALYSIS = {
    "김규탁": {
        "role_summary": "해외물류/출고 승인 흐름에서 주문 검증, 출고 가능 여부 판단, 수출·통관 서류와 예외 승인 근거를 확인하는 역할이 강하게 나타난다.",
        "work_done": [
            "해외물류 방에서 주문번호별 출하 가능 여부, 수입자 연락, 운임·인보이스·CIPL·HS/세번 관련 확인을 처리했다.",
            "핑거세일즈 주문현황에서 품목/단가/수량 오류, 오발송, 회수, 부분 출고 여부를 확인하고 반려 또는 수정 요청을 냈다.",
            "미수·후불·출고보류 건은 영업부의 구두 요청이 아니라 결재·기안·대표 승인 근거가 있어야 한다는 기준을 반복적으로 제시했다.",
            "해외영업과 물류 사이의 공식 요청창구를 NOCO로 모으고, 단순 메신저 요청을 구조화하려는 개선 의견을 냈다.",
        ],
        "rules": [
            "해외 출고 요청은 주문번호, 국가/업체명, 품목, 출하조건, 수금·승인 근거가 함께 있어야 처리한다.",
            "미수 또는 후불 리스크가 큰 주문은 영업부 업무기안, 대표/팀장 결재, 명시 승인 기록을 확인한 뒤 물류가 출고한다.",
            "구두 승인, 단순 전달, '승인 받았다'는 표현만으로는 출고 근거로 보지 않는다.",
            "주문·품목 오류가 발견되면 반려/승인취소 후 담당자가 수정 재등록하거나 승인 절차를 다시 태운다.",
            "해외영업-물류 요청은 NOCO 같은 공식 창구에 기록하고, 긴급 연락도 사후 기록을 남긴다.",
        ],
        "evidence": [
            "2026-04-06 DM: 물류방 요청은 '물건 나가달라, 서류작성해달라, 언제 나갈 수 있냐' 수준의 단순 요청이 많다고 진단.",
            "2026-04-06 DM: 본부장 구두 승인만으로는 의미가 없고, 미수채권방/대표 승인 등 공식 근거가 필요하다고 설명.",
            "2025-08~09 해외물류: OD 주문별 출고 승인 해제, CIPL, EXW/DAP, 수입자 연락, 운임 청구 방식을 확인.",
            "2025-08~09 핑거세일즈 주문현황: 오발송 회수, 품목코드 수정, 부분 출고/익일 발송 여부를 주문번호 기준으로 안내.",
        ],
    },
    "조소연": {
        "role_summary": "경영기획 쪽에서 주문 등록/반려/삭제, 출고 지연 사유, 회계마감·수금 근거, 주문 데이터 품질 기준을 관리하는 역할이 강하다.",
        "work_done": [
            "핑거세일즈 주문현황에서 지연 사유, 매입처 재고, 직발송 가능 여부, 반품/재작성/삭제를 정리했다.",
            "해외물류에서 수금등록, 송금증빙, 주문번호 기재, 출하 근거를 요구하며 기준 미충족 건을 반려했다.",
            "주문서 복사 시 회계마감일, 출고일, 운송사, 운송장번호가 복사되지 않도록 설정하고 확인을 요청했다.",
            "NOCO 사용 원칙, 긴급 건은 유선 더블체크 필요, 기존 채팅방과 신규 공식창구가 병행되는 문제를 설명했다.",
        ],
        "rules": [
            "출고/주문 요청은 주문번호, 고객/국가, 요청 내용, 지연·보류 사유를 간략하고 명확하게 남긴다.",
            "수금·송금증빙·회계 등록이 없으면 출하 근거 부족으로 보며, 담당자에게 근거 보강 후 재요청하게 한다.",
            "주문 수정/반려/삭제 이력은 히스토리로 추적 가능해야 하며, A/S 주문은 장비 및 시리얼번호 같은 필수값을 요구한다.",
            "주문 취소·반품·오출고는 신규 주문, 음수 매출, 회수월 처리 등 회계월 기준으로 정리한다.",
            "일반 요청은 NOCO/공식 채널에 남기고, 긴급 건만 유선/메신저로 더블체크한 뒤 기록한다.",
        ],
        "evidence": [
            "2025-09-15 해외물류: 회계팀 수금등록과 송금증빙이 없어 출하 근거가 부족하니 근거 제출 후 요청하라고 안내.",
            "2025-09-15 해외물류: 주문번호를 기재하고 요청 내용만 간략히 적어 누락을 방지하라고 안내.",
            "2026-03-26 DM: 주문 진행→반려→수정 절차에서 세부 이력 확인 필요성을 제기.",
            "2026-05-08 DM: 긴급 건은 유선 더블체크가 반드시 필요하며, 요청사항/해결은 NOCO로 가야 한다고 설명.",
        ],
    },
    "정재회": {
        "role_summary": "물류/출고 실행 담당으로, 재고 부족 확인, 소모재 전달, 회수·재발송, 쇼핑몰/핑거 주문 등록 보조가 많이 나타난다.",
        "work_done": [
            "핑거세일즈 주문현황에서 주문번호별 재고 부족, 수량·단가 오입력, 전달 대상자와 설치일을 확인했다.",
            "소모재를 특정 팀/담당자에게 전달하고, 담당자가 부재하면 자리에 두거나 추가 확인을 진행했다.",
            "회수 접수, 재발송, 송장번호 공유, 재고 부족 시 부분 출고 여부를 담당자에게 확인했다.",
            "쇼핑몰 주문을 핑거세일즈에 등록하거나, 출고승인취소/반려 권한 문제를 바로 제기했다.",
        ],
        "rules": [
            "재고 부족은 주문번호, 고객명, 품목, 부족 수량, 가능한 출고일을 함께 남기고 부분 출고 여부를 확인한다.",
            "전달 업무는 수령자/팀/장소를 확인하고, 부재 시 어디에 두었는지 기록한다.",
            "오출고·회수·재발송은 택배 회수 접수 여부와 송장번호를 함께 공유한다.",
            "단가/수량/품목 코드 오류는 주문 승인 전 바로 공유하고 수정 또는 반려 가능자를 연결한다.",
            "권한이 없어 반려/승인취소가 안 되는 경우 즉시 시스템 권한 요청을 올린다.",
        ],
        "evidence": [
            "2025-08-28 핑거세일즈 주문현황: 특정 소모재를 TECH팀/신용훈 대리에게 전달할지 확인하고 부재자 자리에 두었다고 기록.",
            "2025-08-28 핑거세일즈 주문현황: 킴스치과기공소 주문에서 단가와 수량이 반대로 입력된 것 같다고 공유.",
            "2025-09-01 핑거세일즈 주문현황: 재고 부족 품목의 가능한 출고일과 나머지 품목 선출고 여부를 확인.",
            "2026-04-15 DM: 출고승인취소/반려 권한이 없다고 시스템 권한 이슈를 보고.",
        ],
    },
    "김채원": {
        "role_summary": "사업전략/영업관리에서 국내 주문 정정, 신규 품목/코드 요청, 온라인샵 정책·고객 리스트 정리, 빠른 출고 요청을 수행했다.",
        "work_done": [
            "핑거세일즈 주문현황에서 주문명, 단가, 유/무상 구분, 담당자, 고객 정보 오류를 발견하면 반려 요청 후 수정 재등록했다.",
            "신규 품목/코드 생성 기안을 올리고, 필요한 담당자에게 생성 요청을 전달했다.",
            "퀵 발송, 부분 발송, 고객 연락처, 주문 사유를 주문번호와 함께 물류/출고 담당자에게 전달했다.",
            "온라인샵의 적립금/예치금/할인코드, 툴 가격 인상 유예기간, 구매 제한 같은 운영 정책을 주문 처리 기준으로 연결했다.",
            "고객사/잠재고객 리스트를 지역, 주소, 연락처 중심으로 정리하고 업무규칙/기능 요구를 AX 쪽으로 넘기는 역할도 보였다.",
        ],
        "rules": [
            "주문 오류는 주문번호와 오류 사유를 명시해 반려 요청하고, 수정 주문을 재등록한다.",
            "신규 품목은 기안→품목 코드 생성 요청→주문 반영 순서로 진행한다.",
            "급한 출고/퀵 요청은 주문번호, 주문명, 품목, 선불/착불, 급한 사유를 함께 남긴다.",
            "온라인샵 정책 변경은 할인코드, 적립금/예치금, 가격 인상 유예, 구매 제한을 주문 입력 기준으로 번역해 반영한다.",
            "업무규칙 수집/AX 기능 논의는 운영 개선 근거로 보되, 단순 개발 요청은 통상업무 규칙에서 분리한다.",
        ],
        "evidence": [
            "2026-01-21 핑거세일즈 주문현황: 단가 오기입으로 주문 반려를 요청하고 수정 주문을 다시 진행.",
            "2026-01-21 핑거세일즈 주문현황: 신규품목 기안을 올리고 품목 코드 생성을 요청.",
            "2026-02-02 핑거세일즈 주문현황: 퀵 발송 요청 시 주문번호, 주문명, 사유, 선불/착불을 함께 전달.",
            "2026-06-04 DM: 예치금/적립금 분류와 할인코드, 툴 가격 인상 유예기간·구매 제한 운영을 확인.",
        ],
    },
    "이미연": {
        "role_summary": "사업전략/영업관리에서 국내 주문 등록·재작성·반려 요청, 온라인샵 주문 보정, 출고팀 더블체크, 인수인계 정리를 담당했다.",
        "work_done": [
            "핑거세일즈 주문현황에서 주문 삭제, 재등록, 신규품목 신청, 임원승인 단계 확인, 반려 요청을 반복 수행했다.",
            "소모재/장비 전달 대상자를 확인하고 물류/정재회에게 전달 요청을 냈다.",
            "온라인샵 주문 자동화 알림을 보고 회사명, 고객명, 주문명, 품목, 적립금/할인 반영 오류를 보정했다.",
            "출고팀과 더블체크하여 출고 완료 여부를 확인하고, 완료건 제목 수정 권한 같은 운영 권한을 요청했다.",
            "신규 담당자 인수인계에서 업무 파일, 예외처리, 주문 처리 순서를 정리하겠다고 했다.",
        ],
        "rules": [
            "국내 주문 오류는 먼저 수정/재등록하고, 필요한 경우 조소연/김규탁에게 삭제·반려·승인 확인을 요청한다.",
            "소모재 전달은 담당자와 수령자, 설치일, 재고 여부를 확인한 뒤 물류 담당자에게 명확히 넘긴다.",
            "온라인샵 주문은 noti-order 알림을 확인하고 회사명/고객명/주문명/품목/결제정책 오류를 사람이 최종 보정한다.",
            "출고 완료 여부는 출고팀과 더블체크하며, 완료 이후 제목/담당자 수정 권한이 필요하면 별도 요청한다.",
            "인수인계는 '어떤 파일을 보는지, 어떤 예외는 어떻게 처리하는지, 주문할 때 1-2-3 순서가 무엇인지'를 문서화한다.",
        ],
        "evidence": [
            "2025-08~09 핑거세일즈 주문현황: 주문 삭제 가능 여부, 신규품목신청, 반려 요청, 임원승인 단계 확인을 수행.",
            "2025-08-28 핑거세일즈 주문현황: 더서울디지털/케이치과기공소 소모재 전달 대상을 정재회에게 안내.",
            "2026-01-09 DM: 출고는 출고쪽과 자신이 함께 더블체크한다고 설명.",
            "2026-01-20 DM: 인수인계 중 특이사항을 정리해 전달하겠다고 응답.",
        ],
    },
}


def load_json(path: Path, default=None):
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def iter_messages():
    for path in (ROOT / "chats").glob("*.json"):
        payload = load_json(path, {})
        source = (payload.get("chat") or {}).get("name") or path.stem
        for message in payload.get("messages") or []:
            yield source, message.get("ts"), message.get("who") or "", clean_text(message.get("text") or "")
    for path in (ROOT / "channels").glob("*.json"):
        payload = load_json(path, {})
        channel = payload.get("channel") or {}
        source = f"{channel.get('team_name') or ''} / {channel.get('channel_name') or ''}".strip()
        for thread in payload.get("threads") or []:
            for message in [thread.get("parent") or {}] + (thread.get("replies") or []):
                yield source, message.get("ts"), message.get("who") or "", clean_text(message.get("text") or "")


def categories_for(source: str, text: str) -> list[str]:
    haystack = f"{source} {text}".lower()
    return [
        category
        for category, keywords in CATEGORY_KEYWORDS.items()
        if any(keyword.lower() in haystack for keyword in keywords)
    ]


def person_stats(messages: list[tuple[str, str, str, str]]) -> dict[str, dict]:
    stats = {}
    for person in TARGETS:
        authored = [(source, ts, who, text) for source, ts, who, text in messages if person in who]
        source_counts = Counter(source for source, *_ in authored)
        category_counts = Counter()
        for source, _, _, text in authored:
            category_counts.update(categories_for(source, text))
        stats[person] = {
            "authored_count": len(authored),
            "top_sources": source_counts.most_common(8),
            "category_counts": category_counts.most_common(),
        }
    return stats


def collection_stats() -> dict:
    chatrooms = load_json(ROOT / "chatrooms.json", {"chats": []}).get("chats") or []
    channels = [c for c in (load_json(ROOT / "channels.json", {"channels": []}).get("channels") or []) if c.get("channel_id")]
    chat_files = list((ROOT / "chats").glob("*.json"))
    channel_files = list((ROOT / "channels").glob("*.json"))

    chat_messages = 0
    for path in chat_files:
        chat_messages += load_json(path, {}).get("message_count") or 0

    channel_messages = 0
    channel_threads = 0
    reply_errors = 0
    completed_channels = []
    for path in channel_files:
        payload = load_json(path, {})
        channel = payload.get("channel") or {}
        completed_channels.append({
            "team": channel.get("team_name") or "",
            "channel": channel.get("channel_name") or "",
            "messages": payload.get("message_count") or 0,
            "threads": payload.get("thread_count") or 0,
            "reply_errors": payload.get("reply_error_count") or 0,
        })
        channel_messages += payload.get("message_count") or 0
        channel_threads += payload.get("thread_count") or 0
        reply_errors += payload.get("reply_error_count") or 0

    status = load_json(ROOT / "collect-status.json", {})
    previous = load_json(Path("output/teams/conversations/_summary.json"), {})
    prev_teams = previous.get("teams") or {}

    return {
        "chatrooms_total": len(chatrooms),
        "chatroom_kinds": Counter(c.get("kind") for c in chatrooms),
        "joined_teams": len({c.get("team_id") for c in channels}),
        "channels_total": len(channels),
        "chat_files": len(chat_files),
        "chat_messages": chat_messages,
        "channel_files": len(channel_files),
        "channel_messages": channel_messages,
        "channel_threads": channel_threads,
        "channel_reply_errors": reply_errors,
        "status_errors": status.get("errors") or [],
        "channels": channels,
        "completed_channels": sorted(completed_channels, key=lambda r: (r["team"], r["channel"])),
        "previous": {
            "extracted_at": previous.get("extracted_at"),
            "channels": prev_teams.get("channels"),
            "threads": prev_teams.get("channel_threads_total"),
            "replies": prev_teams.get("channel_replies_total"),
            "errors": len(prev_teams.get("team_errors") or []),
        },
    }


def person_info(person: str) -> dict:
    return load_json(PEOPLE_DIR / f"{person}_person_info.json", {})


def md_list(items: list[str]) -> list[str]:
    return [f"- {item}" for item in items]


def render_markdown(stats: dict, pstats: dict[str, dict]) -> str:
    lines = [
        "# Teams 전체 수집 및 인물별 업무규칙 분석",
        "",
        f"- 작성일: {RUN_DATE}",
        f"- 기준 계정: 백인식 Brad <ins@doflab.com>",
        f"- 원천 데이터: `{ROOT}`",
        f"- 산출물: `{OUT}`",
        "",
        "## 1. Teams 수집 범위",
        "",
        f"- 채팅방 inventory: {stats['chatrooms_total']}개 "
        f"(DM {stats['chatroom_kinds'].get('oneOnOne', 0)}, group {stats['chatroom_kinds'].get('group', 0)}, meeting {stats['chatroom_kinds'].get('meeting', 0)})",
        f"- 채팅방 히스토리: {stats['chat_files']}개 파일, {stats['chat_messages']:,} 메시지",
        f"- Joined teams: {stats['joined_teams']}개",
        f"- 채널 inventory: {stats['channels_total']}개",
        f"- 채널 히스토리 완료 파일: {stats['channel_files']}개, {stats['channel_threads']:,} threads, {stats['channel_messages']:,} messages",
        f"- 채널 reply 오류 수: {stats['channel_reply_errors']}",
        f"- collect-status 오류 수: {len(stats['status_errors'])}",
        "",
        "### 미수집 채널",
        "",
    ]
    if stats["status_errors"]:
        for error in stats["status_errors"]:
            channel = error.get("channel") or {}
            lines.append(
                f"- {channel.get('team_name')} / {channel.get('channel_name')}: "
                f"{(error.get('error') or '')[:180]}"
            )
    else:
        lines.append("- 없음")
    lines.extend([
        "",
        "### 기존 수집 대비",
        "",
        f"- 기존 `output/teams/conversations` 요약: {stats['previous']['channels']}개 채널, "
        f"{stats['previous']['threads']} threads, {stats['previous']['replies']} replies, "
        f"{stats['previous']['errors']}개 오류",
        "- 기존 night routine 계열 수집은 날짜별 raw/증분 중심이었다. 이번 산출물은 현재 Graph 권한으로 보이는 전체 채팅방과 전체 팀/채널 inventory를 기준으로 처음부터 현재까지의 메시지 파일을 별도로 만들었다.",
        "- 2026-06-08 기존 전체 dump에서 403으로 실패했던 private 채널도 이번에는 같은 inventory에 포함해 재시도하도록 구성했다.",
        "",
        "## 2. 전체 채널 목록",
        "",
        "| team | channel | type | collected |",
        "|---|---|---|---:|",
    ])
    completed_key = {(c["team"], c["channel"]): c for c in stats["completed_channels"]}
    for channel in stats["channels"]:
        team = channel.get("team_name") or ""
        ch = channel.get("channel_name") or ""
        done = completed_key.get((team, ch))
        collected = "" if not done else f"{done['messages']} msg / {done['threads']} th"
        lines.append(f"| {team} | {ch} | {channel.get('membership_type') or ''} | {collected} |")

    lines.extend([
        "",
        "## 3. 분석 기준",
        "",
        "- 업무규칙은 '어떤 요청이 들어오면 어떤 확인/반려/승인/출고/정정 절차로 처리하는가'를 기준으로 보았다.",
        "- 개발 요청, 포털/CRM/AX 기능 구현 요청, 버그 수정 요청은 통상업무 규칙에서 분리했다. 단, 운영 규칙을 설명하기 위해 나온 시스템 요구는 별도 근거로만 반영했다.",
        "",
    ])

    for offset, person in enumerate(TARGETS, start=4):
        info = person_info(person)
        analysis = PERSON_ANALYSIS[person]
        stat = pstats.get(person, {})
        lines.extend([
            f"## {offset}. {person}",
            "",
            f"- 부서/직책: {info.get('dept', '')} / {info.get('position', '')}",
            f"- ERP main_work: {info.get('main_work')}",
            f"- Teams 작성 메시지: {stat.get('authored_count', 0):,}건",
            f"- 주요 출처: {', '.join(f'{name}({count})' for name, count in stat.get('top_sources', [])[:6])}",
            f"- 작성 메시지 카테고리: {', '.join(f'{name}({count})' for name, count in stat.get('category_counts', [])[:6])}",
            "",
            f"요약: {analysis['role_summary']}",
            "",
            "### 수행 업무",
            "",
            *md_list(analysis["work_done"]),
            "",
            "### 업무규칙",
            "",
            *md_list(analysis["rules"]),
            "",
            "### 근거 예시",
            "",
            *md_list(analysis["evidence"]),
            "",
        ])
    return "\n".join(lines)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def render_docx(stats: dict, pstats: dict[str, dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(9.5)

    doc.add_heading("Teams 전체 수집 및 인물별 업무규칙 분석", 0)
    doc.add_paragraph(f"작성일: {RUN_DATE}")
    doc.add_paragraph("기준 계정: 백인식 Brad <ins@doflab.com>")
    doc.add_paragraph(f"원천 데이터: {ROOT}")

    doc.add_heading("1. Teams 수집 범위", 1)
    add_kv_table(doc, [
        ("채팅방 inventory", f"{stats['chatrooms_total']}개 (DM {stats['chatroom_kinds'].get('oneOnOne', 0)}, group {stats['chatroom_kinds'].get('group', 0)}, meeting {stats['chatroom_kinds'].get('meeting', 0)})"),
        ("채팅방 히스토리", f"{stats['chat_files']}개 파일, {stats['chat_messages']:,} 메시지"),
        ("Joined teams", f"{stats['joined_teams']}개"),
        ("채널 inventory", f"{stats['channels_total']}개"),
        ("채널 히스토리", f"{stats['channel_files']}개 파일, {stats['channel_threads']:,} threads, {stats['channel_messages']:,} messages"),
        ("채널 reply 오류", str(stats["channel_reply_errors"])),
        ("collect-status 오류", str(len(stats["status_errors"]))),
        ("기존 전체 dump", f"{stats['previous']['channels']}개 채널, {stats['previous']['threads']} threads, {stats['previous']['replies']} replies, 오류 {stats['previous']['errors']}개"),
    ])
    doc.add_paragraph("기존 night routine 계열 수집은 날짜별 raw/증분 중심이었다. 이번 산출물은 현재 Graph 권한으로 보이는 전체 채팅방과 전체 팀/채널 inventory를 기준으로 별도 full-history raw를 만들었다.")
    doc.add_heading("미수집 채널", 2)
    if stats["status_errors"]:
        for error in stats["status_errors"]:
            channel = error.get("channel") or {}
            doc.add_paragraph(
                f"{channel.get('team_name')} / {channel.get('channel_name')}: "
                f"{(error.get('error') or '')[:220]}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("없음", style="List Bullet")

    doc.add_heading("2. 전체 채널 목록", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "team"
    hdr[1].text = "channel"
    hdr[2].text = "type"
    hdr[3].text = "collected"
    completed_key = {(c["team"], c["channel"]): c for c in stats["completed_channels"]}
    for channel in stats["channels"]:
        team = channel.get("team_name") or ""
        ch = channel.get("channel_name") or ""
        done = completed_key.get((team, ch))
        row = table.add_row().cells
        row[0].text = team
        row[1].text = ch
        row[2].text = channel.get("membership_type") or ""
        row[3].text = "" if not done else f"{done['messages']} msg / {done['threads']} th"

    doc.add_heading("3. 분석 기준", 1)
    add_bullets(doc, [
        "업무규칙은 어떤 요청이 들어오면 어떤 확인/반려/승인/출고/정정 절차로 처리하는가를 기준으로 보았다.",
        "개발 요청, 포털/CRM/AX 기능 구현 요청, 버그 수정 요청은 통상업무 규칙에서 분리했다.",
        "운영 규칙을 설명하기 위해 나온 시스템 요구는 별도 근거로만 반영했다.",
    ])

    for person in TARGETS:
        info = person_info(person)
        analysis = PERSON_ANALYSIS[person]
        stat = pstats.get(person, {})
        doc.add_heading(person, 1)
        add_kv_table(doc, [
            ("부서/직책", f"{info.get('dept', '')} / {info.get('position', '')}"),
            ("ERP main_work", str(info.get("main_work"))),
            ("Teams 작성 메시지", f"{stat.get('authored_count', 0):,}건"),
            ("주요 출처", ", ".join(f"{name}({count})" for name, count in stat.get("top_sources", [])[:6])),
            ("작성 메시지 카테고리", ", ".join(f"{name}({count})" for name, count in stat.get("category_counts", [])[:6])),
        ])
        doc.add_paragraph(f"요약: {analysis['role_summary']}")
        doc.add_heading("수행 업무", 2)
        add_bullets(doc, analysis["work_done"])
        doc.add_heading("업무규칙", 2)
        add_bullets(doc, analysis["rules"])
        doc.add_heading("근거 예시", 2)
        add_bullets(doc, analysis["evidence"])

    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    stats = collection_stats()
    messages = list(iter_messages())
    pstats = person_stats(messages)
    REPORT_MD.write_text(render_markdown(stats, pstats), encoding="utf-8")
    render_docx(stats, pstats)
    print(json.dumps({
        "markdown": str(REPORT_MD),
        "docx": str(REPORT_DOCX),
        "chat_messages": stats["chat_messages"],
        "channel_messages": stats["channel_messages"],
        "channel_files": stats["channel_files"],
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
